"""Compose a standalone academic Word results pack from validated inputs."""

from __future__ import annotations

from collections.abc import Iterable, Sequence
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .model import (
    CoefficientRecord,
    CoefficientTable,
    DescriptiveRecord,
    DescriptiveTable,
    DocumentSpec,
    FigureInput,
    NarrativeMode,
    NarrativeClaim,
    ResultsDocxError,
)
from .word_style import (
    add_body,
    add_caption,
    add_heading,
    add_table_note,
    add_title_block,
    apply_table_geometry,
    configure_document,
    distribute_widths,
    mark_header_row,
    page_break,
    set_cell_text,
    set_picture_alt_text,
    set_row_rule,
    shade_row,
)


TRANSPORT_PACK_LABEL = "TRANSPORT VERIFICATION PACK"
TRANSPORT_TITLE = "Structured statistical output"
TRANSPORT_SUBTITLE = "Machine-recorded values and provenance only"


def compose_document(spec: DocumentSpec, *, author: str) -> tuple[Document, tuple[NarrativeClaim, ...]]:
    document = Document()
    document.core_properties.author = author
    document.core_properties.last_modified_by = author
    transport_only = spec.narrative_mode == "transport-only"
    pack_label = TRANSPORT_PACK_LABEL if transport_only else "EMPIRICAL RESULTS PACK"
    configure_document(document, pack_label=pack_label)

    input_count = 1 + len(spec.coefficient_tables) + len(spec.descriptive_tables) + sum(
        1 + len(figure.source_data) + (1 if figure.source_script else 0) for figure in spec.figures
    )
    add_title_block(
        document,
        title=TRANSPORT_TITLE if transport_only else spec.title,
        subtitle=TRANSPORT_SUBTITLE if transport_only else spec.subtitle,
        run_id=spec.run_id,
        as_of_date=spec.as_of_date,
        input_count=input_count,
        output_identity=author,
        pack_label=pack_label,
    )

    claims: list[NarrativeClaim] = []
    overview_heading = (
        "Transport verification overview"
        if spec.narrative_mode == "transport-only"
        else "Results overview"
    )
    add_heading(document, overview_heading, 1)
    intro_claims = [
        _coefficient_claim(table, spec.coefficient_decimals, spec.narrative_mode, index)
        for index, table in enumerate(spec.coefficient_tables, 1)
    ]
    intro_text = " ".join(claim.text for claim in intro_claims)
    add_body(document, intro_text)
    claims.extend(intro_claims)
    if spec.narrative_mode == "transport-only":
        add_body(
            document,
            "Transport-only mode reproduces values for pipeline verification and does not assess or interpret "
            "the reported model outputs.",
            italic=True,
            muted=True,
        )
    else:
        add_body(
            document,
            "These statements report estimated associations and sampling uncertainty. They do not, by themselves, establish a causal effect.",
            italic=True,
            muted=True,
        )

    if spec.descriptive_tables:
        descriptive_heading = (
            "Recorded descriptive outputs"
            if spec.narrative_mode == "transport-only"
            else "Sample and descriptive statistics"
        )
        add_heading(document, descriptive_heading, 1)
        for table_index, table in enumerate(spec.descriptive_tables, 1):
            claim = _descriptive_claim(table, spec.descriptive_decimals, spec.narrative_mode, table_index)
            add_body(document, claim.text)
            claims.append(claim)
            _add_descriptive_table(
                document,
                table,
                spec.descriptive_decimals,
                spec.narrative_mode,
                table_index,
            )

    if transport_only and spec.descriptive_tables:
        page_break(document)
    model_heading = "Recorded model outputs" if spec.narrative_mode == "transport-only" else "Regression evidence"
    add_heading(document, model_heading, 1)
    for index, table in enumerate(spec.coefficient_tables, 1):
        if index > 1:
            page_break(document)
        _add_coefficient_table(document, table, spec.coefficient_decimals, spec.narrative_mode, index)

    included_figures = embedded_figures(spec)
    if included_figures:
        page_break(document)
        figure_heading = "Transport figures" if spec.narrative_mode == "transport-only" else "Figures"
        add_heading(document, figure_heading, 1)
        for index, figure in enumerate(included_figures, 1):
            _add_figure(document, figure, spec.narrative_mode, index)

    audit_heading = "Transport audit boundary" if spec.narrative_mode == "transport-only" else "Audit boundary"
    add_heading(document, audit_heading, 1)
    boundary = (
        "Displayed statistical values are copied from the hashed machine-readable inputs. Free-form titles, labels, "
        "notes, outcome descriptions, and model descriptions are not used in transport-only output. An explicitly "
        "included image is transported as binary content; its pixels are not semantically audited by this builder. "
        "No manuscript file is edited by this workflow."
        if transport_only
        else "Every displayed number is derived from the machine-readable inputs listed in the companion manifest. "
        "The manifest records file hashes and row-level provenance for generated narrative claims. No manuscript "
        "file is edited by this workflow."
    )
    add_body(document, boundary)
    return document, tuple(claims)


def embedded_figures(spec: DocumentSpec) -> tuple[FigureInput, ...]:
    """Return figures that the selected narrative policy permits in the DOCX."""
    if spec.narrative_mode == "standard":
        return spec.figures
    return tuple(figure for figure in spec.figures if figure.transport_figure)


def _add_descriptive_table(
    document: Document,
    table_spec: DescriptiveTable,
    decimals: int,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> None:
    transport_only = narrative_mode == "transport-only"
    add_caption(document, _descriptive_title(table_spec, narrative_mode, table_index))
    columns = _descriptive_columns(table_spec.records, narrative_mode)
    table = document.add_table(rows=1, cols=len(columns))
    apply_table_geometry(table, distribute_widths(2_260, len(columns) - 1))
    mark_header_row(table.rows[0])
    shade_row(table.rows[0])
    for cell, (_, label) in zip(table.rows[0].cells, columns, strict=True):
        set_cell_text(cell, label, bold=True, size=8.5)
    set_row_rule(table.rows[0], top=True, bottom=True, size=8)

    for record in table_spec.records:
        row = table.add_row()
        values = [_descriptive_value(record, key, decimals) for key, _ in columns]
        for index, (cell, value) in enumerate(zip(row.cells, values, strict=True)):
            set_cell_text(
                cell,
                value,
                align=WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                size=8.5,
            )
    set_row_rule(table.rows[-1], bottom=True, size=8)
    note = (
        "Notes: Values are mechanically formatted from the recorded numeric fields. N is the recorded non-missing "
        "observation count. Free-form labels and notes are excluded in transport-only mode."
        if transport_only
        else table_spec.note
        or "Notes: N is the non-missing observation count. Statistics are reproduced from the tidy descriptive-statistics input."
    )
    add_table_note(document, note)
    add_table_note(
        document,
        f"Audit source SHA-256: {table_spec.source.sha256}. Full provenance is in RESULTS_DOCX_MANIFEST.md."
        if transport_only
        else f"Audit source: {table_spec.source.path.name} (SHA-256 {table_spec.source.sha256[:12]}...); full provenance is in RESULTS_DOCX_MANIFEST.md.",
    )


def _add_coefficient_table(
    document: Document,
    table_spec: CoefficientTable,
    decimals: int,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> None:
    transport_only = narrative_mode == "transport-only"
    add_caption(document, _coefficient_title(table_spec, narrative_mode, table_index))
    panels = _first_seen(record.panel for record in table_spec.records if record.row_type == "coef") or [""]
    for panel_index, panel in enumerate(panels):
        if len(panels) > 1 or panel:
            add_table_note(document, f"Panel {panel or panel_index + 1}", keep_with_next=True)
        panel_records = [
            record
            for record in table_spec.records
            if record.panel == panel or (record.row_type != "coef" and not record.panel)
        ]
        _add_coefficient_panel(document, panel_records, decimals, narrative_mode)

    if transport_only:
        note = (
            "Notes: Standard errors are in parentheses. Star symbols are mechanical display transformations of the "
            "recorded two-sided p-values at 0.01, 0.05, and 0.10. Free-form labels, model descriptions, outcome "
            "descriptions, specification text, and notes are excluded in transport-only mode."
        )
    elif table_spec.note:
        note = table_spec.note
    else:
        note = (
            "Notes: Standard errors are in parentheses. ***, **, and * indicate two-sided p-values below 0.01, 0.05, and 0.10, respectively."
        )
    add_table_note(document, note)
    add_table_note(
        document,
        f"Audit source SHA-256: {table_spec.source.sha256}. Full row-level provenance is in RESULTS_DOCX_MANIFEST.md."
        if transport_only
        else f"Audit source: {table_spec.source.path.name} (SHA-256 {table_spec.source.sha256[:12]}...); full row-level provenance is in RESULTS_DOCX_MANIFEST.md.",
    )


def _add_coefficient_panel(
    document: Document,
    records: Sequence[CoefficientRecord],
    decimals: int,
    narrative_mode: NarrativeMode,
) -> None:
    transport_only = narrative_mode == "transport-only"
    coef_records = [record for record in records if record.row_type == "coef"]
    model_order = _first_seen(record.model_id for record in coef_records)
    term_order = _first_seen(record.term for record in coef_records)
    model_labels = {
        model: model if transport_only else _single_model_value(coef_records, model, "model_label") or model
        for model in model_order
    }

    explicit_rows = [] if transport_only else _first_seen(record.term for record in records if record.row_type != "coef")
    metadata_rows = _metadata_footer_rows(coef_records, model_order, explicit_rows, narrative_mode)
    total_rows = 1 + (2 * len(term_order)) + len(explicit_rows) + len(metadata_rows)
    table = document.add_table(rows=total_rows, cols=1 + len(model_order))
    apply_table_geometry(table, distribute_widths(2_900, len(model_order)))

    header = table.rows[0]
    mark_header_row(header)
    shade_row(header)
    set_cell_text(header.cells[0], "Variable", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    for index, model in enumerate(model_order, 1):
        set_cell_text(header.cells[index], model_labels[model], bold=True, size=9)
    set_row_rule(header, top=True, bottom=True, size=8)

    row_index = 1
    for term in term_order:
        estimate_row = table.rows[row_index]
        se_row = table.rows[row_index + 1]
        term_records = {record.model_id: record for record in coef_records if record.term == term}
        term_label = term if transport_only else next(record.term_label for record in coef_records if record.term == term)
        set_cell_text(estimate_row.cells[0], term_label, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        set_cell_text(se_row.cells[0], "", align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        for model_index, model in enumerate(model_order, 1):
            record = term_records.get(model)
            estimate = "" if record is None or record.estimate is None else _format_estimate(record.estimate, record.p_value, decimals)
            standard_error = "" if record is None or record.std_error is None else f"({_format_number(record.std_error, decimals)})"
            set_cell_text(estimate_row.cells[model_index], estimate, size=9)
            set_cell_text(se_row.cells[model_index], standard_error, italic=True, size=9)
        row_index += 2

    for term in explicit_rows:
        row = table.rows[row_index]
        matching = {record.model_id: record for record in records if record.row_type != "coef" and record.term == term}
        first = next(iter(matching.values()))
        set_cell_text(row.cells[0], first.term_label, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        for model_index, model in enumerate(model_order, 1):
            record = matching.get(model)
            value = _footer_record_value(record, decimals) if record else ""
            set_cell_text(row.cells[model_index], value, size=9)
        row_index += 1

    for label, values in metadata_rows:
        row = table.rows[row_index]
        set_cell_text(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        for model_index, model in enumerate(model_order, 1):
            set_cell_text(row.cells[model_index], values.get(model, ""), size=9)
        row_index += 1

    set_row_rule(table.rows[-1], bottom=True, size=8)


def _add_figure(
    document: Document,
    figure: FigureInput,
    narrative_mode: NarrativeMode,
    figure_index: int,
) -> None:
    transport_only = narrative_mode == "transport-only"
    title = f"Transport figure F{figure_index}" if transport_only else figure.title
    alt_text = (
        f"Transported figure F{figure_index}. Image pixels are not semantically audited."
        if transport_only
        else figure.alt_text
    )
    add_caption(document, title)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Inches(0.03)
    paragraph.paragraph_format.space_after = Inches(0.06)
    run = paragraph.add_run()
    run.add_picture(str(figure.path), width=Inches(5.8))
    set_picture_alt_text(run, title=title, description=alt_text)
    if transport_only:
        add_table_note(
            document,
            "Transport boundary: this image was explicitly opted in as binary evidence. The builder verifies its "
            "file hash and provenance inputs but does not semantically audit labels, annotations, or visual claims "
            "inside the pixels.",
        )
    elif figure.note:
        add_table_note(document, figure.note)
    add_table_note(
        document,
        f"Audit image SHA-256: {figure.source.sha256}. Full provenance hashes are in RESULTS_DOCX_MANIFEST.md."
        if transport_only
        else "Audit source: image "
        + figure.path.name
        + "; source data "
        + ", ".join(item.path.name for item in figure.source_data)
        + ". Full hashes are in RESULTS_DOCX_MANIFEST.md.",
    )


def _coefficient_claim(
    table: CoefficientTable,
    decimals: int,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> NarrativeClaim:
    coefficients = [record for record in table.records if record.row_type == "coef"]
    non_intercepts = [record for record in coefficients if record.term.lower() not in {"(intercept)", "intercept", "_cons"}]
    candidates = non_intercepts or coefficients
    if table.primary_term:
        candidates = [record for record in candidates if record.term == table.primary_term]
        if not candidates:
            raise ResultsDocxError(f"{table.source.path}: primary_term {table.primary_term!r} not found")
    if table.primary_model:
        candidates = [record for record in candidates if record.model_id == table.primary_model]
        if not candidates:
            raise ResultsDocxError(f"{table.source.path}: primary_model {table.primary_model!r} not found for primary term")
    record = candidates[-1]
    assert record.estimate is not None and record.std_error is not None and record.p_value is not None
    n_text = f", N={record.nobs:,}" if record.nobs is not None else ""
    if narrative_mode == "transport-only":
        text = (
            f"Coefficient output C{table_index}: term_id={record.term}; model_id={record.model_id}; "
            f"estimate={_format_number(record.estimate, decimals)}; "
            f"standard_error={_format_number(record.std_error, decimals)}; "
            f"p_value={_format_p(record.p_value)}{n_text}."
        )
    else:
        outcome = record.dependent_variable or "the reported outcome"
        direction = "positive" if record.estimate > 0 else "negative" if record.estimate < 0 else "zero"
        significance = _significance_phrase(record.p_value)
        text = (
            f"{table.title} reports a {direction} estimated association between {record.term_label} and {outcome} "
            f"in {record.model_label}: coefficient {_format_number(record.estimate, decimals)}, standard error "
            f"{_format_number(record.std_error, decimals)}, p={_format_p(record.p_value)}{n_text}; the estimate is {significance}."
        )
    return NarrativeClaim(
        text=text,
        source_path=table.source.path,
        source_row=record.source_row,
        selectors={"term": record.term, "model_id": record.model_id, "panel": record.panel},
        values=(
            {
                "estimate": record.estimate,
                "std.error": record.std_error,
                "p.value": record.p_value,
                "nobs": record.nobs,
            }
            if narrative_mode == "transport-only"
            else {
                "estimate": record.estimate,
                "std.error": record.std_error,
                "p.value": record.p_value,
                "nobs": record.nobs,
                "dependent_variable": record.dependent_variable,
            }
        ),
    )


def _descriptive_claim(
    table: DescriptiveTable,
    decimals: int,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> NarrativeClaim:
    record = table.records[0]
    metrics: list[str] = []
    values: dict[str, str | int | float | None] = {}
    if record.n is not None:
        metrics.append(f"N={record.n:,}")
        values["n"] = record.n
    if record.mean is not None:
        metrics.append(f"mean={_format_number(record.mean, decimals)}")
        values["mean"] = record.mean
    if record.sd is not None:
        metrics.append(f"SD={_format_number(record.sd, decimals)}")
        values["sd"] = record.sd
    if narrative_mode == "transport-only":
        text = (
            f"Descriptive output D{table_index}: variable_id={record.variable}; "
            + "; ".join(metrics)
            + "."
        )
    else:
        text = f"{table.title} summarizes the analysis sample. For {record.variable_label}, " + ", ".join(metrics) + "."
    return NarrativeClaim(
        text=text,
        source_path=table.source.path,
        source_row=record.source_row,
        selectors=(
            {"variable": record.variable}
            if narrative_mode == "transport-only"
            else {"variable": record.variable, "sample": record.sample}
        ),
        values=values,
    )


def _descriptive_columns(
    records: Sequence[DescriptiveRecord],
    narrative_mode: NarrativeMode,
) -> list[tuple[str, str]]:
    identifier_field = "variable" if narrative_mode == "transport-only" else "variable_label"
    candidates = [
        (identifier_field, "Variable ID" if narrative_mode == "transport-only" else "Variable"),
        ("n", "N"),
        ("mean", "Mean"),
        ("sd", "SD"),
        ("p25", "P25"),
        ("p50", "Median"),
        ("p75", "P75"),
        ("minimum", "Min"),
        ("maximum", "Max"),
    ]
    return [
        item
        for item in candidates
        if item[0] in {"variable", "variable_label"}
        or any(getattr(record, item[0]) is not None for record in records)
    ]


def _descriptive_value(record: DescriptiveRecord, key: str, decimals: int) -> str:
    value = getattr(record, key)
    if value is None:
        return ""
    if key in {"variable", "variable_label"}:
        return str(value)
    if key == "n":
        return f"{value:,}"
    return _format_number(float(value), decimals)


def _metadata_footer_rows(
    records: Sequence[CoefficientRecord],
    models: Sequence[str],
    explicit_terms: Sequence[str],
    narrative_mode: NarrativeMode,
) -> list[tuple[str, dict[str, str]]]:
    definitions = (
        ("Controls", "controls", {"controls"}, lambda value: str(value)),
        ("Fixed effects", "fixed_effects", {"fe", "fixedeffects", "fixed_effects"}, lambda value: str(value)),
        ("SE clustered by", "cluster", {"cluster", "clustering", "seclusteredby"}, lambda value: str(value)),
        ("Observations", "nobs", {"n", "nobs", "observations"}, lambda value: f"{int(value):,}"),
        (
            "Adjusted R-squared",
            "adj_r_squared",
            {"adjr2", "adjrsquared", "adjustedrsquared"},
            lambda value: _format_number(float(value), 3),
        ),
    )
    if narrative_mode == "transport-only":
        definitions = tuple(item for item in definitions if item[1] in {"nobs", "adj_r_squared"})
    normalized_explicit = {_normalize_footer_term(term) for term in explicit_terms}
    rows: list[tuple[str, dict[str, str]]] = []
    for label, field, aliases, formatter in definitions:
        if normalized_explicit & aliases:
            continue
        values: dict[str, str] = {}
        for model in models:
            value = _single_model_value(records, model, field)
            if value not in (None, ""):
                values[model] = formatter(value)
        if values:
            rows.append((label, values))
    return rows


def _normalize_footer_term(term: str) -> str:
    return "".join(character for character in term.lower() if character.isalnum() or character == "_")


def _coefficient_title(
    table: CoefficientTable,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> str:
    return f"Coefficient output C{table_index}" if narrative_mode == "transport-only" else table.title


def _descriptive_title(
    table: DescriptiveTable,
    narrative_mode: NarrativeMode,
    table_index: int,
) -> str:
    return f"Descriptive output D{table_index}" if narrative_mode == "transport-only" else table.title


def _single_model_value(records: Sequence[CoefficientRecord], model: str, field: str) -> object:
    values = [getattr(record, field) for record in records if record.model_id == model and getattr(record, field) not in (None, "")]
    return values[0] if values else None


def _footer_record_value(record: CoefficientRecord, decimals: int) -> str:
    if record.value_text:
        return record.value_text
    if record.estimate is not None:
        return _format_number(record.estimate, decimals)
    if record.nobs is not None:
        return f"{record.nobs:,}"
    if record.adj_r_squared is not None:
        return _format_number(record.adj_r_squared, decimals)
    return ""


def _format_estimate(estimate: float, p_value: float | None, decimals: int) -> str:
    stars = ""
    if p_value is not None:
        if p_value < 0.01:
            stars = "***"
        elif p_value < 0.05:
            stars = "**"
        elif p_value < 0.10:
            stars = "*"
    return f"{_format_number(estimate, decimals)}{stars}"


def _format_number(value: float, decimals: int) -> str:
    rounded = round(value, decimals)
    if rounded == 0:
        rounded = 0.0
    return f"{rounded:.{decimals}f}"


def _format_p(value: float) -> str:
    return "<0.001" if value < 0.001 else f"{value:.3f}"


def _significance_phrase(p_value: float) -> str:
    if p_value < 0.01:
        return "statistically distinguishable from zero at the 1% level"
    if p_value < 0.05:
        return "statistically distinguishable from zero at the 5% level"
    if p_value < 0.10:
        return "statistically distinguishable from zero at the 10% level"
    return "not statistically distinguishable from zero at the 10% level"


def _first_seen(values: Iterable[str]) -> list[str]:
    return list(dict.fromkeys(values))
