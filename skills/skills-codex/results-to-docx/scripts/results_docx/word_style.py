"""Word-native styling and OOXML geometry for the academic results pack."""

from __future__ import annotations

from collections.abc import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


PAGE_WIDTH_DXA = 12_240
PAGE_HEIGHT_DXA = 15_840
CONTENT_WIDTH_DXA = 9_360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

SERIF = "Times New Roman"
NAVY = RGBColor(46, 116, 181)
DARK_NAVY = RGBColor(31, 77, 120)
INK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"


def configure_document(document: Document, *, pack_label: str = "EMPIRICAL RESULTS PACK") -> None:
    """Apply the resolved standard_business_brief token map and academic override."""
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(11)
    _set_style_font_xml(normal.element, SERIF)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = document.styles["Title"]
    title.font.name = SERIF
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.underline = False
    title.font.color.rgb = INK
    _set_style_font_xml(title.element, SERIF)
    _remove_style_paragraph_borders(title.element)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = SERIF
    subtitle.font.size = Pt(12)
    subtitle.font.italic = True
    subtitle.font.color.rgb = MUTED
    _set_style_font_xml(subtitle.element, SERIF)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = SERIF
    caption.font.size = Pt(11)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = INK
    _set_style_font_xml(caption.element, SERIF)
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    kicker = _ensure_paragraph_style(document, "Results Pack Kicker")
    kicker.font.name = SERIF
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = NAVY
    _set_style_font_xml(kicker.element, SERIF)
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(8)

    table_note = _ensure_paragraph_style(document, "Academic Table Note")
    table_note.font.name = SERIF
    table_note.font.size = Pt(9)
    table_note.font.italic = True
    table_note.font.color.rgb = MUTED
    _set_style_font_xml(table_note.element, SERIF)
    table_note.paragraph_format.space_before = Pt(4)
    table_note.paragraph_format.space_after = Pt(4)

    heading_tokens = {
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, NAVY, 12, 6),
        "Heading 3": (12, DARK_NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = SERIF
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        _set_style_font_xml(style.element, SERIF)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for section in document.sections:
        _add_running_header(section, pack_label)
        _add_running_footer(section)


def add_title_block(
    document: Document,
    *,
    title: str,
    subtitle: str,
    run_id: str,
    as_of_date: str,
    input_count: int,
    output_identity: str,
    pack_label: str = "EMPIRICAL RESULTS PACK",
) -> None:
    kicker = document.add_paragraph(pack_label, style="Results Pack Kicker")

    title_p = document.add_paragraph(title, style="Title")

    subtitle_p = document.add_paragraph(subtitle, style="Subtitle")

    metadata = (
        ("Run ID", run_id),
        ("As of", as_of_date),
        ("Audited inputs", str(input_count)),
        ("Output identity", output_identity),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        set_run(paragraph.add_run(f"{label}: "), size=9.5, bold=True, color=INK)
        set_run(paragraph.add_run(value), size=9.5, color=INK)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    _set_paragraph_bottom_border(rule, color="2E74B5", size=12)


def add_heading(document: Document, text: str, level: int = 1) -> Paragraph:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(document: Document, text: str, *, italic: bool = False, muted: bool = False) -> Paragraph:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run(run, size=11, italic=italic, color=MUTED if muted else INK)
    return paragraph


def add_caption(document: Document, text: str) -> Paragraph:
    return document.add_paragraph(text, style="Caption")


def add_table_note(document: Document, text: str, *, keep_with_next: bool = False) -> Paragraph:
    paragraph = document.add_paragraph(text, style="Academic Table Note")
    paragraph.paragraph_format.keep_with_next = keep_with_next
    return paragraph


def set_run(
    run: Run,
    *,
    size: float,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor = INK,
) -> None:
    run.font.name = SERIF
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), SERIF)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), SERIF)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    size: float = 9,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.widow_control = False
    paragraph.clear()
    set_run(paragraph.add_run(text), size=size, bold=bold, italic=italic)


def apply_table_geometry(table: Table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must total {CONTENT_WIDTH_DXA} DXA, got {sum(widths_dxa)}")
    if len(widths_dxa) != len(table.columns):
        raise ValueError("Column width count must match table column count")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    _set_or_replace_dxa(tbl_pr, "tblW", CONTENT_WIDTH_DXA)
    _set_or_replace_dxa(tbl_pr, "tblInd", TABLE_INDENT_DXA)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    _set_table_cell_margins(tbl_pr)
    _clear_table_borders(tbl_pr)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            _set_or_replace_dxa(tc_pr, "tcW", widths_dxa[index])
            cell.width = Inches(widths_dxa[index] / 1440)


def distribute_widths(first_column_dxa: int, other_columns: int) -> list[int]:
    if other_columns < 1:
        return [CONTENT_WIDTH_DXA]
    remainder = CONTENT_WIDTH_DXA - first_column_dxa
    base, extra = divmod(remainder, other_columns)
    return [first_column_dxa] + [base + (1 if index < extra else 0) for index in range(other_columns)]


def mark_header_row(row: object) -> None:
    tr_pr = row._tr.get_or_add_trPr()  # type: ignore[attr-defined]
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def set_row_rule(row: object, *, top: bool = False, bottom: bool = False, size: int = 8) -> None:
    for cell in row.cells:  # type: ignore[attr-defined]
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge, enabled in (("top", top), ("bottom", bottom)):
            if not enabled:
                continue
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), "000000")


def shade_row(row: object, fill: str = LIGHT_GRAY) -> None:
    for cell in row.cells:  # type: ignore[attr-defined]
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)


def set_picture_alt_text(run: Run, *, title: str, description: str) -> None:
    for node in run._element.xpath(".//wp:docPr | .//pic:cNvPr"):
        node.set("name", title)
        node.set("title", title)
        node.set("descr", description)


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _add_running_header(section: object, pack_label: str) -> None:
    paragraph = section.header.paragraphs[0]  # type: ignore[attr-defined]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.clear()
    set_run(paragraph.add_run(pack_label), size=8.5, bold=True, color=MUTED)
    _set_paragraph_bottom_border(paragraph, color="D9E2F3", size=4)


def _add_running_footer(section: object) -> None:
    paragraph = section.footer.paragraphs[0]  # type: ignore[attr-defined]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.clear()
    set_run(paragraph.add_run("ARIS audit-ready output  |  "), size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), SERIF)
    fonts.set(qn("w:hAnsi"), SERIF)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "595959")
    properties.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    properties.append(size)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _set_style_font_xml(style_element: object, font_name: str) -> None:
    r_pr = style_element.get_or_add_rPr()  # type: ignore[attr-defined]
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)


def _ensure_paragraph_style(document: Document, name: str) -> object:
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _remove_style_paragraph_borders(style_element: object) -> None:
    p_pr = style_element.find(qn("w:pPr"))  # type: ignore[attr-defined]
    if p_pr is None:
        return
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def _set_or_replace_dxa(parent: object, local_name: str, value: int) -> None:
    node = parent.find(qn(f"w:{local_name}"))  # type: ignore[attr-defined]
    if node is None:
        node = OxmlElement(f"w:{local_name}")
        parent.append(node)  # type: ignore[attr-defined]
    node.set(qn("w:w"), str(value))
    node.set(qn("w:type"), "dxa")


def _set_table_cell_margins(tbl_pr: object) -> None:
    margins = tbl_pr.find(qn("w:tblCellMar"))  # type: ignore[attr-defined]
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)  # type: ignore[attr-defined]
    for edge, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _clear_table_borders(tbl_pr: object) -> None:
    borders = tbl_pr.find(qn("w:tblBorders"))  # type: ignore[attr-defined]
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)  # type: ignore[attr-defined]
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def _set_paragraph_bottom_border(paragraph: Paragraph, *, color: str, size: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
